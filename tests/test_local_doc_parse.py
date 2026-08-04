"""로컬 문서 파싱 단위 테스트입니다."""

from __future__ import annotations

import io
from unittest.mock import patch

import pytest

from app.utils.local_doc_parse import (
    file_extension,
    parse_docx_bytes,
    parse_local_document_sync,
    parse_plain_bytes,
    pdf_bytes_to_parsed,
    render_table_png,
)


def test_file_extension():
    assert file_extension("a.HWP") == "hwp"
    assert file_extension("x.y.docx") == "docx"


def test_parse_plain_bytes():
    parsed = parse_plain_bytes("안녕하세요".encode("utf-8"), "a.txt")
    assert parsed is not None
    assert "안녕" in parsed.contents[0]["content"]
    assert parsed.total_pages == 1


def test_parse_doc_unsupported():
    assert parse_local_document_sync(b"x", "legacy.doc", for_chat=False) is None


def test_render_table_png():
    png = render_table_png([["A", "B"], ["1", "2"]], title="테스트")
    assert png and png[:8] == b"\x89PNG\r\n\x1a\n"


def test_parse_docx_with_table():
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("서문 문단")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "이름"
    table.rows[0].cells[1].text = "점수"
    table.rows[1].cells[0].text = "홍길동"
    table.rows[1].cells[1].text = "100"
    doc.save(buf)

    parsed = parse_docx_bytes(buf.getvalue(), for_chat=True)
    assert parsed is not None
    text = parsed.contents[0]["content"]
    assert "서문 문단" in text
    assert "홍길동" in text
    assert parsed.page_images, "표 이미지가 있어야 합니다"


def test_hwp_uses_rhwp_then_pdf(monkeypatch):
    """HWP는 rhwp→PDF→pypdfium 경로를 탄다 (바이너리는 mock)."""

    # 최소 유효 PDF (빈 페이지 1개) — pypdfium이 열 수 있어야 함
    # 테스트를 단순화하기 위해 pdf_bytes_to_parsed를 패치
    fake_pdf = b"%PDF-1.4 fake"

    with patch("app.utils.local_doc_parse.hwp_to_pdf_bytes", return_value=fake_pdf) as mock_hwp:
        with patch(
            "app.utils.local_doc_parse.pdf_bytes_to_parsed",
            return_value=__import__("app.utils.local_doc_parse", fromlist=["ParsedDocument"]).ParsedDocument(
                contents=[{"content": "조항 텍스트", "page": 1}],
                metadata={"source": "hwp"},
                total_pages=1,
                page_images=[],
            ),
        ):
            parsed = parse_local_document_sync(b"hwp-bytes", "규정.hwp", for_chat=False)
            assert parsed is not None
            assert parsed.contents[0]["content"] == "조항 텍스트"
            mock_hwp.assert_called_once()


def test_build_user_content_with_text_and_images():
    from app.utils.chat_attachments import ResolvedAttachment, build_user_content

    uri = "data:image/png;base64,AAAA"
    content = build_user_content(
        "요약해줘",
        [
            ResolvedAttachment(
                attachment_id="1",
                file_name="a.pdf",
                file_type="application/pdf",
                kind="document",
                object_key="k",
                text="본문",
                data_uris=[uri],
            )
        ],
    )
    assert isinstance(content, list)
    assert content[0]["type"] == "image_url"
    assert content[-1]["type"] == "text"
    assert "본문" in content[-1]["text"]


@pytest.mark.skipif(
    not __import__("os").environ.get("RHWP_BIN")
    and not __import__("shutil").which("rhwp"),
    reason="rhwp 바이너리 없음",
)
def test_rhwp_smoke_real_hwp():
    from pathlib import Path

    from app.utils.local_doc_parse import find_rhwp_bin

    assert find_rhwp_bin()
    samples = list(Path("resources/regulations").glob("1-0-1*.hwp"))
    if not samples:
        pytest.skip("샘플 HWP 없음")
    data = samples[0].read_bytes()
    parsed = parse_local_document_sync(data, samples[0].name, for_chat=True, max_image_pages=1)
    assert parsed is not None
    assert parsed.total_pages >= 1
    assert parsed.contents or parsed.page_images
