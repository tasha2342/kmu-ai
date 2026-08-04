"""로컬 문서 파싱 파사드입니다.

외부 Doc Parser HTTP API를 대체합니다.
- HWP/HWPX: rhwp CLI → PDF → pypdfium2 (텍스트 + 선택적 페이지 이미지)
- PDF: pypdfium2
- DOCX: python-docx (+ 챗용 표 PNG)
- TXT/MD: UTF-8 등 디코딩
- DOC: 미지원

학칙 인제스트(hwp_extractor/pyhwp)와는 별개입니다.
"""

from __future__ import annotations

import asyncio
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Any, Optional

from app.utils.logger import get_logger

logger = get_logger("local_doc_parse", log_dir="logs")

RHWP_VERSION = "v0.8.2"
"""Docker에 설치하는 rhwp 릴리즈 태그 (문서·Dockerfile과 맞출 것)"""

DEFAULT_RENDER_SCALE = 1.5
DEFAULT_MAX_IMAGE_PAGES = 4


@dataclass
class ParsedDocument:
    """로컬 파싱 결과. RAG는 contents만, 챗은 page_images도 사용합니다."""

    contents: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    total_pages: int = 0
    page_images: list[dict[str, Any]] = field(default_factory=list)
    """챗용. 항목: {page, mime, data_uri}"""


def file_extension(file_name: str) -> str:
    return PurePosixPath(file_name or "").suffix.lstrip(".").lower()


def find_rhwp_bin() -> Optional[str]:
    """RHWP_BIN 또는 PATH에서 rhwp 실행 파일을 찾습니다."""

    env_path = (os.environ.get("RHWP_BIN") or "").strip()
    if env_path and Path(env_path).is_file() and os.access(env_path, os.X_OK):
        return env_path
    return shutil.which("rhwp")


def _image_data_uri_png(png_bytes: bytes) -> str:
    from app.utils.chat_attachments import image_data_uri

    return image_data_uri("image/png", png_bytes)


def pdf_bytes_to_parsed(
    pdf_data: bytes,
    *,
    for_chat: bool,
    max_image_pages: int = DEFAULT_MAX_IMAGE_PAGES,
    scale: float = DEFAULT_RENDER_SCALE,
    source: str = "pdf",
) -> Optional[ParsedDocument]:
    """PDF 바이트에서 페이지 텍스트(및 선택 이미지)를 뽑습니다."""

    if not pdf_data:
        return None

    import pypdfium2 as pdfium

    contents: list[dict[str, Any]] = []
    page_images: list[dict[str, Any]] = []

    pdf = pdfium.PdfDocument(pdf_data)
    try:
        page_count = len(pdf)
        if page_count <= 0:
            return None

        image_limit = min(page_count, max(1, max_image_pages)) if for_chat else 0

        for index in range(page_count):
            page = pdf[index]
            try:
                textpage = page.get_textpage()
                try:
                    text = (textpage.get_text_bounded() or "").strip()
                finally:
                    textpage.close()
                if text:
                    contents.append({"content": text, "page": index + 1})

                if for_chat and index < image_limit:
                    bitmap = page.render(scale=scale)
                    pil_image = bitmap.to_pil()
                    buf = BytesIO()
                    pil_image.save(buf, format="PNG", optimize=True)
                    page_images.append({
                        "page": index + 1,
                        "mime": "image/png",
                        "data_uri": _image_data_uri_png(buf.getvalue()),
                    })
            finally:
                page.close()
    finally:
        pdf.close()

    if not contents and not page_images:
        return None

    return ParsedDocument(
        contents=contents,
        metadata={"source": source, "parser": "pypdfium2"},
        total_pages=page_count,
        page_images=page_images,
    )


def _noto_font_args() -> list[str]:
    """컨테이너에 설치된 Noto CJK 계열을 rhwp fallback으로 넘깁니다."""

    args: list[str] = []
    # Debian fonts-noto-cjk 패키지 경로
    search_roots = [
        Path("/usr/share/fonts/opentype/noto"),
        Path("/usr/share/fonts/truetype/noto"),
    ]
    for root in search_roots:
        if root.is_dir():
            args.extend(["--font-path", str(root)])
            break
    # 패밀리 이름은 Debian/Ubuntu Noto CJK 패키지 기준
    args.extend([
        "--fallback-sans", "Noto Sans CJK KR",
        "--fallback-serif", "Noto Serif CJK KR",
        "--fallback-mono", "Noto Sans Mono CJK KR",
    ])
    return args


def hwp_to_pdf_bytes(
    file_data: bytes,
    file_name: str,
    *,
    rhwp_bin: Optional[str] = None,
    timeout_seconds: int = 120,
) -> Optional[bytes]:
    """rhwp export-pdf로 HWP/HWPX를 PDF 바이트로 변환합니다."""

    binary = rhwp_bin or find_rhwp_bin()
    if not binary:
        logger.error("rhwp 실행 파일을 찾을 수 없습니다. RHWP_BIN 또는 PATH를 확인하세요.")
        return None

    ext = file_extension(file_name) or "hwp"
    with tempfile.TemporaryDirectory(prefix="rhwp_") as work_dir:
        work = Path(work_dir)
        src = work / f"input.{ext}"
        out_pdf = work / "out.pdf"
        src.write_bytes(file_data)

        cmd = [binary, "export-pdf", str(src), "-o", str(out_pdf), *_noto_font_args()]
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=max(1, timeout_seconds),
                check=False,
            )
        except subprocess.TimeoutExpired:
            logger.error(f"rhwp export-pdf 타임아웃 ({timeout_seconds}s): {file_name}")
            return None
        except FileNotFoundError:
            logger.error(f"rhwp 실행 실패 (파일 없음): {binary}")
            return None

        if result.returncode != 0 or not out_pdf.is_file():
            stderr = (result.stderr or result.stdout or "").strip()
            logger.error(
                f"rhwp export-pdf 실패 (code={result.returncode}, file={file_name}): {stderr[:500]}"
            )
            return None

        return out_pdf.read_bytes()


def _docx_font(size: int = 16):
    from PIL import ImageFont

    candidates = (
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansKR-Regular.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    )
    for name in candidates:
        if Path(name).exists():
            try:
                return ImageFont.truetype(name, size)
            except Exception:
                continue
    return ImageFont.load_default()


def render_table_png(rows: list[list[str]], title: str = "표") -> Optional[bytes]:
    """표 행을 PNG 바이트로 렌더링합니다 (챗 멀티모달용)."""

    from PIL import Image, ImageDraw

    if not rows:
        return None

    font = _docx_font(16)
    title_font = _docx_font(18)
    width = max((len(r) for r in rows), default=1)
    norm = [list(r) + [""] * (width - len(r)) for r in rows]

    tmp = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(tmp)
    col_w = [40] * width
    row_h: list[int] = []
    for row in norm:
        rh = 24
        for i, cell in enumerate(row):
            text = str(cell)[:80] or " "
            bbox = draw.textbbox((0, 0), text, font=font)
            tw = min(bbox[2] - bbox[0] + 16, 240)
            th = bbox[3] - bbox[1] + 12
            col_w[i] = max(col_w[i], tw)
            rh = max(rh, th)
        row_h.append(rh)

    title_h = 36
    pad = 16
    table_w = sum(col_w) + 1
    table_h = sum(row_h) + 1
    img_w = max(table_w + pad * 2, 480)
    img_h = title_h + table_h + pad * 2

    img = Image.new("RGB", (img_w, img_h), "white")
    draw = ImageDraw.Draw(img)
    draw.text((pad, 8), title[:80], fill="black", font=title_font)

    y = title_h
    for ri, row in enumerate(norm):
        x = pad
        for ci, cell in enumerate(row):
            w, h = col_w[ci], row_h[ri]
            fill = "#e8eef7" if ri == 0 else "white"
            draw.rectangle([x, y, x + w, y + h], outline="#333333", fill=fill)
            draw.text((x + 6, y + 4), str(cell)[:40] or " ", fill="black", font=font)
            x += w
        y += row_h[ri]

    buf = BytesIO()
    img.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


def parse_docx_bytes(file_data: bytes, *, for_chat: bool) -> Optional[ParsedDocument]:
    """DOCX에서 문단·표 텍스트(및 선택적 표 이미지)를 뽑습니다."""

    if not file_data:
        return None

    from docx import Document

    try:
        doc = Document(BytesIO(file_data))
    except Exception:
        logger.exception("DOCX 열기에 실패했습니다.")
        return None

    parts: list[str] = []
    page_images: list[dict[str, Any]] = []
    table_index = 0

    # 문단과 표를 문서 순서대로 최대한 반영 (body 순회)
    body = doc.element.body
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, doc)
            text = (para.text or "").strip()
            if text:
                parts.append(text)
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            table_index += 1
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([(cell.text or "").strip().replace("\n", " ") for cell in row.cells])
            md_lines = [f"[표 {table_index}]"]
            if rows:
                header = rows[0]
                md_lines.append("| " + " | ".join(header) + " |")
                md_lines.append("| " + " | ".join("---" for _ in header) + " |")
                for row in rows[1:]:
                    md_lines.append("| " + " | ".join(row) + " |")
            parts.append("\n".join(md_lines))
            if for_chat and rows:
                png = render_table_png(rows, title=f"표 {table_index}")
                if png:
                    page_images.append({
                        "page": table_index,
                        "mime": "image/png",
                        "data_uri": _image_data_uri_png(png),
                    })

    text = "\n\n".join(parts).strip()
    if not text and not page_images:
        return None

    return ParsedDocument(
        contents=[{"content": text, "page": 1}] if text else [],
        metadata={"source": "docx", "parser": "python-docx", "table_count": table_index},
        total_pages=1,
        page_images=page_images,
    )


def parse_plain_bytes(file_data: bytes, file_name: str) -> Optional[ParsedDocument]:
    """TXT/MD 디코딩."""

    if not file_data:
        return None
    text = None
    for encoding in ("utf-8", "utf-8-sig", "cp949", "euc-kr", "latin-1"):
        try:
            text = file_data.decode(encoding).replace("\x00", "").strip()
            break
        except UnicodeDecodeError:
            continue
    if not text:
        return None
    return ParsedDocument(
        contents=[{"content": text, "page": 1}],
        metadata={"source": file_extension(file_name), "parser": "plain"},
        total_pages=1,
        page_images=[],
    )


def parse_local_document_sync(
    file_data: bytes,
    file_name: str,
    *,
    for_chat: bool = False,
    max_image_pages: int = DEFAULT_MAX_IMAGE_PAGES,
    rhwp_timeout_seconds: int = 120,
) -> Optional[ParsedDocument]:
    """동기 로컬 파싱. 호출 측에서 to_thread로 감싸는 것을 권장합니다."""

    ext = file_extension(file_name)

    if ext == "doc":
        logger.warning("legacy .doc 형식은 로컬 파서가 지원하지 않습니다.")
        return None

    if ext in {"txt", "md"}:
        return parse_plain_bytes(file_data, file_name)

    if ext == "docx":
        return parse_docx_bytes(file_data, for_chat=for_chat)

    if ext == "pdf":
        return pdf_bytes_to_parsed(
            file_data,
            for_chat=for_chat,
            max_image_pages=max_image_pages,
            source="pdf",
        )

    if ext in {"hwp", "hwpx"}:
        pdf_data = hwp_to_pdf_bytes(
            file_data,
            file_name,
            timeout_seconds=rhwp_timeout_seconds,
        )
        if not pdf_data:
            return None
        parsed = pdf_bytes_to_parsed(
            pdf_data,
            for_chat=for_chat,
            max_image_pages=max_image_pages,
            source=ext,
        )
        if parsed:
            parsed.metadata["parser"] = "rhwp+pypdfium2"
            parsed.metadata["rhwp_version"] = RHWP_VERSION
        return parsed

    logger.warning(f"지원하지 않는 확장자입니다: {ext}")
    return None


async def parse_local_document(
    file_data: bytes,
    file_name: str,
    *,
    for_chat: bool = False,
    max_image_pages: Optional[int] = None,
    rhwp_timeout_seconds: Optional[int] = None,
) -> Optional[ParsedDocument]:
    """비동기 로컬 파싱 진입점."""

    try:
        from app.config import config

        pages = max_image_pages
        if pages is None:
            pages = getattr(config.chatbot, "attachment_image_max_pages", DEFAULT_MAX_IMAGE_PAGES)
        timeout = rhwp_timeout_seconds
        if timeout is None:
            timeout = getattr(config.chatbot, "rhwp_timeout_seconds", 120)
    except Exception:
        pages = max_image_pages or DEFAULT_MAX_IMAGE_PAGES
        timeout = rhwp_timeout_seconds or 120

    return await asyncio.to_thread(
        parse_local_document_sync,
        file_data,
        file_name,
        for_chat=for_chat,
        max_image_pages=pages,
        rhwp_timeout_seconds=timeout,
    )


def parsed_to_legacy_dict(parsed: Optional[ParsedDocument]) -> Optional[dict[str, Any]]:
    """기존 Doc Parser 응답 shape으로 변환합니다 (RAG 호환)."""

    if not parsed:
        return None
    return {
        "contents": parsed.contents,
        "metadata": parsed.metadata,
        "total_pages": parsed.total_pages,
    }
